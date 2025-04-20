import os
import io
import pdfplumber  
import docx
from flask import jsonify
from openai import OpenAI
import re
import json
import logging
from datetime import datetime
import uuid
import time
import requests
from dotenv import load_dotenv

# Load environment variables
load_dotenv(".env.local")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Get API keys from environment variables with fallbacks
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    logger.warning("OPENAI_API_KEY not found in environment variables")
    # No fallback for API keys - they should be in environment variables

# Initialize OpenAI client only if API key is available
openai_client = None
if OPENAI_API_KEY:
    openai_client = OpenAI(api_key=OPENAI_API_KEY)

# Get Groq API key from environment
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
if not GROQ_API_KEY:
    logger.warning("GROQ_API_KEY not found in environment variables")
    # No fallback for API keys - they should be in environment variables

GROQ_API_URL = os.environ.get("GROQ_API_URL", "https://api.groq.com/openai/v1/chat/completions")

def clean_response(response_text):
    """
    Remove any thinking tags, internal reasoning, or other system artifacts
    from the response before sending it to the user
    """
    # Remove <Thinking> tags 
    cleaned = re.sub(r'<Thinking>.*?</Thinking>', '', response_text, flags=re.DOTALL)
    
    thinking_phrases = [
        r"Let me think about this\.?",
        r"Let's think through this\.?",
        r"Thinking through this\.?",
        r"I need to reason through this\.?",
        r"Let me analyze this\.?",
        r"Let me break this down\.?",
        r"Let's analyze this step by step\.?",
        r"Let me work through this\.?",
        r"I'll think about this\.?",
        r"Thinking:.*?\n",
        r"Think:.*?\n",
        r"Internal reasoning:.*?\n",
        r"Step \d+:.*?\n"
    ]
    
    for phrase in thinking_phrases:
        cleaned = re.sub(phrase, '', cleaned, flags=re.IGNORECASE)
    
    # Fix markdown formatting
    cleaned = re.sub(r'\*\* (.*?) \*\*', r'**\1**', cleaned)
    cleaned = re.sub(r'^\*([^\s])', r'* \1', cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r'^-([^\s])', r'- \1', cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r'^(\d+)\.([^\s])', r'\1. \2', cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r'^(#+)([^\s])', r'\1 \2', cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    
    return cleaned.strip()

def extract_text(source, file_type):
    """
    Extract text from various file formats
    Accepts either a file path (string) or a file-like object
    """
    try:
        if isinstance(source, str):
            f = open(source, "rb")
            close_after = True
        else:
            f = source
            close_after = False

        text = ""
        if file_type == "pdf":
            try:
                with pdfplumber.open(f) as pdf:
                    logger.info(f"PDF opened successfully. Pages: {len(pdf.pages)}")
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
                            logger.info(f"Extracted {len(page_text)} characters from page {i+1}")
                        else:
                            logger.info(f"No text found on page {i+1}")
            except Exception as e:
                logger.error(f"PDF extraction error: {str(e)}")
                return None
        elif file_type == "docx":
            try:
                f.seek(0)
                doc = docx.Document(f)
                logger.info(f"DOCX opened successfully. Paragraphs: {len(doc.paragraphs)}")
                for i, para in enumerate(doc.paragraphs):
                    if para.text:
                        text += para.text + "\n"
                        if i < 5 or i % 50 == 0:
                            logger.info(f"Paragraph {i+1}: {para.text[:50]}...")
            except Exception as e:
                logger.error(f"DOCX extraction error: {str(e)}")
                return None
        elif file_type == "txt":
            try:
                f.seek(0)
                text = f.read().decode("utf-8")
                logger.info(f"TXT file read successfully with UTF-8 encoding. Size: {len(text)} characters")
            except UnicodeDecodeError:
                logger.info("UTF-8 decoding failed, trying latin-1 encoding")
                try:
                    f.seek(0)
                    text = f.read().decode("latin-1")
                    logger.info(f"TXT file read successfully with latin-1 encoding. Size: {len(text)} characters")
                except Exception as e:
                    logger.error(f"TXT encoding error: {str(e)}")
                    return None
            except Exception as e:
                logger.error(f"TXT extraction error: {str(e)}")
                return None
        if close_after:
            f.close()
        if not text or text.strip() == "":
            logger.warning("No text found in the file.")
            return None
        logger.info(f"Successfully extracted {len(text)} characters from the source.")
        return text
    except Exception as e:
        logger.error(f"Extraction error: {str(e)}")
        return None

def extract_financial_data(text):
    """
    Extract key financial metrics and ratios from the text using regex patterns
    Returns structured data for use in document analysis
    """
    logger.info("Extracting financial data from document text")
    extracted_data = {}
    
    # Extract financial metrics
    metrics = {}
    
    revenue_match = re.search(r"Total Revenue.*?\$([0-9,.]+)M", text, re.IGNORECASE)
    if revenue_match:
        metrics["revenue"] = revenue_match.group(1)
        logger.info(f"Extracted revenue: {revenue_match.group(1)}")
        
    net_income_match = re.search(r"Net Income.*?\$([0-9,.]+)M", text, re.IGNORECASE)
    if net_income_match:
        metrics["net_income"] = net_income_match.group(1)
        logger.info(f"Extracted net income: {net_income_match.group(1)}")
        
    eps_match = re.search(r"Earnings Per Share.*?\$([0-9.]+)", text, re.IGNORECASE)
    if eps_match:
        metrics["eps"] = eps_match.group(1)
        logger.info(f"Extracted EPS: {eps_match.group(1)}")
        
    ebitda_match = re.search(r"EBITDA.*?\$([0-9,.]+)M", text, re.IGNORECASE)
    if ebitda_match:
        metrics["ebitda"] = ebitda_match.group(1)
        logger.info(f"Extracted EBITDA: {ebitda_match.group(1)}")
    
    extracted_data["metrics"] = metrics
    
    ratios = {}
    ratio_patterns = [
        (r"Gross Margin:?\s*([0-9.]+)%", "gross_margin"),
        (r"Operating Margin:?\s*([0-9.]+)%", "operating_margin"),
        (r"Net Profit Margin:?\s*([0-9.]+)%", "net_profit_margin"),
        (r"Return on Assets.*?([0-9.]+)%", "roa"),
        (r"Return on Equity.*?([0-9.]+)%", "roe")
    ]
    
    for pattern, key in ratio_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            ratios[key] = match.group(1)
            logger.info(f"Extracted {key}: {match.group(1)}")
            
    extracted_data["ratios"] = ratios
    
    segments = []
    segment_section_match = re.search(r"Business Segments(.*?)(?:Geographic Distribution|\d+\.\s)", text, re.DOTALL)
    if segment_section_match:
        segment_text = segment_section_match.group(1)
        segment_matches = re.findall(r"(\d+\.\s*([^:]+):[^%]*?Revenue Contribution:\s*([0-9]+)%)", segment_text, re.DOTALL)
        for match in segment_matches:
            segments.append({
                "name": match[1].strip(),
                "revenue_contribution": match[2].strip() + "%"
            })
            logger.info(f"Extracted segment: {match[1].strip()} - {match[2].strip()}%")
    
    extracted_data["segments"] = segments
    
    logger.info(f"Extraction complete. Found {len(metrics)} metrics, {len(ratios)} ratios, and {len(segments)} segments")
    return extracted_data

def analyze_financial_content(text, user_comment=""):
    """
    Analyze financial document content using LLM
    Includes internal reasoning instructions to prevent showing thought process
    """
    try:
        if not text:
            return "No text to analyze."
        
        logger.info(f"Starting analysis of document with {len(text)} characters")
        if user_comment:
            logger.info(f"User provided comment/instructions: {user_comment[:100]}...")
        
        # Detect if this is a financial document
        financial_keywords = ["balance sheet", "income statement", "cash flow", "financial report", 
                             "revenue", "profit margin", "finance", "money", "cash", "ebitda", "shareholders' equity"]
        
        is_financial = any(keyword in text.lower() for keyword in financial_keywords)
        
        if is_financial:
            logger.info("Document identified as financial report")
        else:
            logger.info("Document identified as general text")
        
        # Extracting financial data if it's a financial document
        extracted_data = {}
        if is_financial:
            extracted_data = extract_financial_data(text)
            metrics_str = json.dumps(extracted_data.get("metrics", {}), indent=2)
            ratios_str = json.dumps(extracted_data.get("ratios", {}), indent=2)
            segments_str = json.dumps(extracted_data.get("segments", []), indent=2)
        
        system_message = """You are a financial analyst specializing in detailed financial report analysis.

IMPORTANT: You must use internal reasoning to analyze the document and formulate your response.
This means you should:
1. Analyze the document content
2. Identify key financial metrics and trends
3. Reason through the best way to present your analysis
4. Provide only your final, polished analysis

FORMATTING INSTRUCTIONS:
- Use proper markdown formatting for your response
- For bold text, use **bold text** format (not ** bold text **)
- For bullet points, use proper markdown: "- Point 1" with a space after the dash
- For numbered lists, use: "1. First item" with a space after the number
- For headings, use: "## Heading" with a space after the #
- Ensure there are no spaces between the asterisks and the text for bold/italic formatting

DO NOT share your internal reasoning process with the user. They should only see your final analysis.
DO NOT use <Thinking> or <Thinking> tags in your response.
DO NOT include phrases like "Let me think about this" or "Analyzing this step by step".
DO NOT number your reasoning steps or include any meta-commentary about your thinking process.
"""
        
        # Create appropriate prompt based on document type
        if is_financial:
            user_message = f"""
            Analyze this financial document and provide a comprehensive analysis with the following structure:

            1. TITLE: Create a clear title for the analysis
            2. EXECUTIVE SUMMARY: 2-3 sentences summarizing the overall financial health
            3. OVERALL PERFORMANCE: Analyze revenue, profitability, and cash flow trends
            4. KEY FINANCIAL METRICS: Present the most important metrics with their values and significance
            5. SEGMENT ANALYSIS: Analyze performance across business segments and geographies if mentioned
            6. FINANCIAL HEALTH INDICATORS: Analyze profitability, liquidity, and solvency ratios
            7. FUTURE OUTLOOK: Summarize forecasts and management's forward-looking statements
            8. RISK ASSESSMENT: Identify key risks and mitigation strategies mentioned
            9. CONCLUSION: Provide a balanced conclusion about the company's financial position

            Format your analysis with clear section headings, bullet points for key insights, and concise paragraphs.
            Make sure your analysis is factual, balanced, and based only on information in the document.
            
            USER COMMENT: {user_comment}

            EXTRACTED FINANCIAL METRICS:
            {metrics_str if is_financial else "No metrics automatically extracted."}

            EXTRACTED FINANCIAL RATIOS:
            {ratios_str if is_financial else "No ratios automatically extracted."}

            EXTRACTED BUSINESS SEGMENTS:
            {segments_str if is_financial else "No segments automatically extracted."}

            DOCUMENT TEXT (EXCERPT):
            {text[:6000]}
            """
        else:
            user_message = f"""
            Analyze this document and provide key insights. Focus on main points, important information, and key takeaways.
            Format your response with clear headings, bullet points for key insights, and concise paragraphs.
            
            USER COMMENT: {user_comment}
            
            Text:
            {text[:4000]}
            """

        logger.info("Sending document for analysis to LLM")
        
        # Try using OpenAI first
        try:
            logger.info("Using OpenAI client for document analysis")
            if not openai_client:
                return "Error: OpenAI API key not configured. Please set the OPENAI_API_KEY environment variable."
                
            response = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": user_message}
                ],
                temperature=0.4,
                max_tokens=1500 if is_financial else 800
            )
            logger.info("Analysis completed successfully with OpenAI")
            response_text = response.choices[0].message.content
            
            # Clean the response to remove any thinking tags
            cleaned_response = clean_response(response_text)
            return cleaned_response
            
        except Exception as e:
            logger.error(f"OpenAI API Error: {str(e)}")
            
            try:
                logger.info("Falling back to Groq API with gemma-7b-it")
                if not GROQ_API_KEY:
                    return "Error: Groq API key not configured. Please set the GROQ_API_KEY environment variable."
                    
                headers = {
                    "Authorization": f"Bearer {GROQ_API_KEY}",
                    "Content-Type": "application/json"
                }
                payload = {
                    "model": "gemma-7b-it",  
                    "messages": [
                        {"role": "system", "content": system_message},
                        {"role": "user", "content": user_message}
                    ],
                    "max_tokens": 1500 if is_financial else 800,
                    "temperature": 0.4
                }
                response = requests.post(GROQ_API_URL, headers=headers, json=payload)
                if response.status_code == 200:
                    logger.info("Analysis completed successfully with Groq API")
                    response_text = response.json()['choices'][0]['message']['content'].strip()
                    
                    cleaned_response = clean_response(response_text)
                    return cleaned_response
                    
                else:
                    error_msg = f"Analysis Error: Groq API - {response.status_code} - {response.text}"
                    logger.error(error_msg)
                    return error_msg
            except Exception as e2:
                error_msg = f"Analysis Error: {str(e)} (OpenAI) and {str(e2)} (Groq)"
                logger.error(error_msg)
                return error_msg

    except Exception as e:
        error_msg = f"AI Processing Error: {str(e)}"
        logger.error(error_msg)
        return f"❌ AI Processing Error: {str(e)}"

def handle_file_upload(request, fs, document_collection, save_to_history_func):
    """
    Handle file upload requests
    This function is imported and used in app.py
    """
    try:
        logger.info("📂 Upload request received")
        logger.info(f"Files in request: {list(request.files.keys())}")
        
        if "file" not in request.files:
            logger.error("❌ No file received in request")
            return jsonify({"error": "No file uploaded"}), 400

        file = request.files["file"]
        logger.info(f"📄 File received: {file.filename}, Size: {file.content_length} bytes")
        if file.filename == "":
            logger.error("❌ Empty filename received")
            return jsonify({"error": "No selected file"}), 400

        user_comment = request.form.get('comment', '')
        if user_comment:
            logger.info(f"📝 User comment received: {user_comment[:100]}...")
            
        user_id = request.form.get('user_id', 'anonymous')
            
        # Ensure the uploaded file is an accepted format
        file_extension = file.filename.split(".")[-1].lower()
        allowed_extensions = {"pdf", "docx", "txt"}
        
        if file_extension not in allowed_extensions:
            logger.error(f"❌ Invalid file type: {file_extension}")
            return jsonify({"error": f"Invalid file type: {file_extension}. Allowed types: pdf, docx, txt"}), 400
            
        # Check file size
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
        if request.content_length and request.content_length > MAX_FILE_SIZE:
            logger.error(f"❌ File too large: {request.content_length} bytes")
            return jsonify({"error": "File is too large. Maximum file size is 10MB."}), 400
            
        file_bytes = file.read()
        gridfs_id = fs.put(file_bytes, filename=file.filename)
        logger.info(f"✅ File stored in MongoDB GridFS with ID: {gridfs_id}")

        file_obj = io.BytesIO(file_bytes)
        extracted_text = extract_text(file_obj, file_extension)
        
        if not extracted_text or extracted_text.strip() == "":
            logger.error("❌ No text found in the file")
            return jsonify({"error": "No text could be extracted from the file"}), 400
            
        document_collection.insert_one({
            "filename": file.filename,
            "gridfs_id": str(gridfs_id),
            "text": extracted_text,
            "upload_date": datetime.utcnow().isoformat(),
            "user_id": user_id
        })
        logger.info(f"✅ Document text stored in database for future reference")
            
        logger.info("🧠 Analyzing text with enhanced LLM analysis...")
        analysis = analyze_financial_content(extracted_text, user_comment)
        logger.info("✅ Analysis complete")
        
        # Save to history
        chat_id = save_to_history_func(
            f"[FILE UPLOAD] {file.filename}" + (f": {user_comment}" if user_comment else ""), 
            analysis, 
            user_id=user_id
        )
        
        return jsonify({
            "gridfs_id": str(gridfs_id),
            "filename": file.filename,
            "analysis": analysis,
            "textLength": len(extracted_text),
            "chatId": chat_id
        })
        
    except Exception as e:
        logger.error(f"❌ Upload Error: {str(e)}")
        return jsonify({"error": f"Upload processing failed: {str(e)}"}), 500

if __name__ == "__main__":
    logger.info("This module provides document processing functionality for the main app.")
    logger.info("It is not intended to be run directly.")